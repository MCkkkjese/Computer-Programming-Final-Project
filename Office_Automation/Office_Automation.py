#coding=utf-8
flag = False
while not flag:
    try:
        import os
        import tkinter as tk
        from tkinter import messagebox
        import ttkbootstrap as ttk
        from PIL import Image, ImageTk
        # import pandas as pd
        # import numpy as np
        import docx
        # import openpyxl
        import matplotlib
        from matplotlib import pyplot as plt
        import yfinance as yf
        import time
        from datetime import datetime, timedelta
        # import selenium
        from selenium import webdriver
        # import pyautogui
        flag = True

    except ImportError :
        print("ERROR : Essential components missed. Automatically installing...")
        from Extension_Modules import install
        install.main()
        os.system("PAUSE")

from Extension_Modules import file_directory as fd
from Extension_Modules import get_username as gu
from Extension_Modules import send_email as se
matplotlib.rc("font", family="Microsoft JhengHei")
print("Application booting up successfully.")

global rect_stock, ticker_list, path_1, path_2, outFile_1, outFile_2
rect_stock = []
ticker_list = []

path_1 = fd.path_function("/rect_stock.dat")
path_2 = fd.path_function("/email_info.dat")
outFile_1 = open(path_1, "r") if os.path.exists(path_1) else open(path_1, "w")
outFile_2 = open(path_2, "r") if os.path.exists(path_2) else open(path_2, "w")

rect_stock = outFile_1.readlines()
for i in range(len(rect_stock)):
    # print(i)
    rect_stock[i] = str(rect_stock[i]).rstrip('\n')

def check_ticker_exists(ticker):
    try:
        stock = yf.Ticker(ticker)
        info = stock.info
        # print(info)  
        if 'regularMarketPrice' in info or 'previousClose' in info:
            return True
        
        else:
            return False
        
    except Exception as e:
        print(f"Error: {e}")  
        return False

def status_output_func(index):
    status_output.config(state="normal")
    status_output.delete("1.0", "end")
    status_output.insert("end", index)
    status_output.config(state="disabled")

def rect_stock_output_func(index):
    rect_stock_output.config(state="normal")
    rect_stock_output.delete("1.0", "end")
    rect_stock_output.insert("end", index)
    rect_stock_output.config(state="disabled")

def apply():
    symbol = stc_sym.get()
    flag = check_ticker_exists(symbol)
    # print(flag)
    if flag:
        ticker_list.append(symbol)
        status_output_func("Stock symbol {} has been added to the pre-operation list.\n".format(symbol))
        rect_stock_output_func("\n".join(ticker_list))

    else:
        status_output_func("The {} does not exist in Yahoo Finance.\n".format(symbol))
        messagebox.showerror("Error", "The stock symbol does not exist in Yahoo Finance.")

    stc_sym.set('')

def appl_all():
    global ticker_list
    # print(rect_stock)
    ticker_list = rect_stock
    status_output_func(f"{len(ticker_list)} stock symbles have been added to the pre-operation list.\n")
    print(ticker_list)

def stock_list():
    if len(ticker_list) == 0:
        messagebox.showerror("Error", "No stock symbol has been added.")
        status_output_func("No stock symbol has been added.\n")
    
    else:
        status_output_func("Stock List : \n" + "\n".join(ticker_list))

def PTC_func(data, ticker, doc):
    plt.figure(figsize=(16, 9))
    plt.plot(data['Close'], label='Close Price')
    plt.title(f'{ticker} - 價格走勢圖')
    plt.xlabel('日期')
    plt.ylabel('價格')
    plt.legend()
    plt.savefig(fd.path_function(f"/{ticker} - Price Trend Chart.png"))
    
    doc.add_heading(f'{ticker} - 價格走勢圖', level=1)
    doc.add_picture((fd.path_function(f"/{ticker} - Price Trend Chart.png")), width=docx.shared.Inches(6), height=docx.shared.Inches(4))
    os.remove(fd.path_function(f"/{ticker} - Price Trend Chart.png"))

def VC_func(data, ticker, doc):
    plt.figure(figsize=(16, 9))
    plt.plot(data['Volume'], label='Volume')
    plt.title(f'{ticker} - 成交量圖表')
    plt.xlabel('日期')
    plt.ylabel('成交量')
    plt.legend()
    plt.savefig(fd.path_function(f"/{ticker} - Volume Chart.png"))
    
    doc.add_heading(f'{ticker} - 成交量圖表', level=1)
    doc.add_picture((fd.path_function(f"/{ticker} - Volume Chart.png")), width=docx.shared.Inches(6), height=docx.shared.Inches(4))
    os.remove(fd.path_function(f"/{ticker} - Volume Chart.png"))

def K_func(data, ticker, doc):
    plt.figure(figsize=(16, 9))
    plt.plot(data['High'], label='High Price')
    plt.plot(data['Low'], label='Low Price')
    plt.plot(data['Open'], label='Open Price')
    plt.plot(data['Close'], label='Close Price')
    plt.title(f'{ticker} - K線圖')
    plt.xlabel('日期')
    plt.ylabel('價格')
    plt.legend()
    plt.savefig(fd.path_function(f"/{ticker} - K Line.png"))

    doc.add_heading(f'{ticker} - K線圖', level=1)
    doc.add_picture((fd.path_function(f"/{ticker} - K Line.png")), width=docx.shared.Inches(6), height=docx.shared.Inches(4))
    os.remove(fd.path_function(f"/{ticker} - K Line.png"))

def MAC_func(data, ticker, MAC_days, doc):
    plt.figure(figsize=(16, 9))
    plt.plot(data['Close'], label='Close Price')
    plt.plot(data['Close'].rolling(MAC_days).mean(), label='MAC')
    plt.title(f'{ticker} - 移動平均線圖')
    plt.xlabel('日期')
    plt.ylabel('價格')
    plt.legend()
    plt.savefig(fd.path_function(f"/{ticker} - Moving Average Chart.png"))

    doc.add_heading(f'{ticker} - 移動平均線圖', level=1)
    doc.add_picture((fd.path_function(f"/{ticker} - Moving Average Chart.png")), width=docx.shared.Inches(6), height=docx.shared.Inches(4))
    os.remove(fd.path_function(f"/{ticker} - Moving Average Chart.png"))

def caculate_RSI(data, window):
    delta = data['Close'].diff()
    gain = (delta.where(delta > 0, 0)).rolling(window=window).mean()
    loss = (-delta.where(delta < 0, 0)).rolling(window=window).mean()
    RS = gain / loss
    RSI = 100 - (100 / (1 + RS))
    return RSI

def RSI_func(data, ticker, days, doc):
    plt.figure(figsize=(16, 9))
    plt.plot(data['Close'], label='Close Price')
    plt.plot(caculate_RSI(data, window=days), label=f'RSI - {days} Days')
    plt.title(f'{ticker} - 相對強弱指數圖表')
    plt.xlabel('日期')
    plt.ylabel('價格')
    plt.legend()
    plt.savefig(fd.path_function(f"/{ticker} - RSI Chart - {days} Days.png"))

    doc.add_heading(f'{ticker} - 相對強弱指數圖表 - {days} Days', level=1)
    doc.add_picture((fd.path_function(f"/{ticker} - RSI Chart - {days} Days.png")), width=docx.shared.Inches(6), height=docx.shared.Inches(4))
    os.remove(fd.path_function(f"/{ticker} - RSI Chart - {days} Days.png"))

def BOL_func(data, ticker, doc):
    data['20_MA'] = data['Close'].rolling(window=20).mean()
    data['20_STD'] = data['Close'].rolling(window=20).std()
    data['Upper Band'] = data['20_MA'] + (data['20_STD'] * 2)
    data['Lower Band'] = data['20_MA'] - (data['20_STD'] * 2)

    plt.figure(figsize=(16, 9))
    plt.plot(data['Close'], label='Close Price')
    plt.plot(data['20_MA'], label='20 Days Moving Average')
    plt.plot(data['Upper Band'], label='Upper Band')
    plt.plot(data['Lower Band'], label='Lower Band')
    plt.title(f'{ticker} - 布林線圖')
    plt.xlabel('日期')
    plt.ylabel('價格')
    plt.legend()
    plt.savefig(fd.path_function(f"/{ticker} - Bollinger Band Chart.png"))

    doc.add_heading(f'{ticker} - 布林線圖', level=1)
    doc.add_picture((fd.path_function(f"/{ticker} - Bollinger Band Chart.png")), width=docx.shared.Inches(6), height=docx.shared.Inches(4))
    os.remove(fd.path_function(f"/{ticker} - Bollinger Band Chart.png"))

# def email_func():
#     pass

def save_to_default():
    SMTP_ = SMTP.get()
    SMTP_ = SMTP_.split('/')
    print(SMTP_)
    _SMTP_ = SMTP_[0]
    TCP_ = int(SMTP_[1])
    email_add_ = email_add.get()
    password_ = password.get()
    path = fd.path_function("/email_info.dat")
    outFile = open(path, "w")
    outFile.write(f"{_SMTP_}/{TCP_}\n{email_add_}\n{password_}")
    outFile.close()

def load_default():
    path = fd.path_function("/email_info.dat")
    inFile = open(path, "r")
    email_info = inFile.readlines()
    SMTP.set(email_info[0].rstrip('\n'))
    email_add.set(email_info[1].rstrip('\n'))
    password.set(email_info[2].rstrip('\n'))

def application():
    status_output_func("Application is running...")
    if len(ticker_list) == 0:
        messagebox.showerror("Error", "No stock symbol has been added.")
        status_output_func("No stock symbol has been added.\n")
    
    else:
        messagebox.showinfo("Information", "Application is running...\nPlease wait until the process is completed.")
        outFile_1 = open(path_1, "w")
        outFile_1.write("\n".join(ticker_list))
        outFile_1.close()
        status_output_func("Stock symbols have been saved.\n")

    time_now = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())
    # print(time_now)
    username = gu.main()

    date_from_ = date_from.get()
    date_to_ = date_to.get()
    PTC_ = PTC.get()
    VC_ = VC.get()
    K_ = K.get()
    MAC_ = MAC.get()
    RSI_ = RSI.get()
    BOL_ = BOL.get()
    email_func_ = email_func.get()

    date_from_ = list(date_from_.split('/'))
    date_to_ = list(date_to_.split('/'))
    try:   
        MAC_days = datetime(int(date_to_[0]), int(date_to_[1]), int(date_to_[2])) - datetime(int(date_from_[0]), int(date_from_[1]), int(date_from_[2]))
    
    except Exception as e:
        print(f"Error: {e}")
        status_output_func("Error: Invalid date format.\n")
        messagebox.showerror("Error", "Invalid date format.\n Please enter the date in the format of YYYY/MM/DD.")
        return
    
    date_from_ = '-'.join(date_from_)
    date_to_ = '-'.join(date_to_)
    print(MAC_days)

    for i in range(len(ticker_list)):
        ticker = ticker_list[i]
        data = yf.download(ticker, start=date_from_, end=date_to_)
        path_3 = "C:\\Users\\{}\\Downloads\\{} - Stock Data {}.docx".format(username, ticker, time_now)
        path_4 = "C:\\Users\\{}\\Downloads\\{} - Stock Data {}.xlsx".format(username, ticker, time_now)
        doc = docx.Document()
        doc.add_heading("Stock Data Analysis Report", 0)
        doc.add_heading("NTTU - Office Automation Final Project - Stock Data Analysis Software", level=9)
        doc.add_heading("Powered by Yahoo Finance - Developed by Liyue-Wei on GitHub, copy right reserved.", level=9)
        doc.add_heading("股票代號: {}".format(ticker), level=1)
        doc.add_paragraph("紀錄日期: {}".format(time.strftime("%Y/%m/%d %H:%M:%S", time.localtime())))
        doc.add_paragraph("日期區間: {} ~ {}".format(date_from_.replace('-', '/'), date_to_.replace('-', '/')))

        if PTC_:
            print("Price Trend Chart")
            PTC_func(data, ticker, doc)

        if VC_:
            print("Volume Chart")
            VC_func(data, ticker, doc)

        if K_:
            print("K Line")
            K_func(data, ticker, doc)
        
        if MAC_:
            print("Moving Average Chart")
            MAC_func(data, ticker, MAC_days, doc)

        if RSI_:
            print("RSI Chart")
            RSI_func(data, ticker, 7, doc)
            RSI_func(data, ticker, 14, doc)
            RSI_func(data, ticker, 21, doc)
            RSI_func(data, ticker, 28, doc)

        if BOL_:
            print("Bollinger Band Chart")
            BOL_func(data, ticker, doc)

        doc.save(path_3)

        # plt.figure(figsize=(14, 7))
        # plt.plot(data['Close'], label='Close Price')
        # plt.title(f'{ticker} 價格走勢圖')
        # plt.xlabel('日期')
        # plt.ylabel('價格')
        # plt.legend()
        # plt.show()

        if email_func_:
            SMTP_ = SMTP.get()
            SMTP_ = SMTP_.split('/')
            print(SMTP_)
            _SMTP_ = SMTP_[0]
            TCP_ = int(SMTP_[1])
            email_add_ = email_add.get()
            password_ = password.get()
            # email_func(_SMTP_, TCP_, email_add_, password_, path_3)
            se.main(email_add_, email_add_, "{} - Stock Data Analysis Report".format(ticker), "{} - Stock Data Analysis Report".format(ticker), _SMTP_, TCP_, password_, path_3)

        # time.sleep(1)
        # status_output_func("Stock data : {} has been saved.\n".format(ticker))
        # time.sleep(1)

    status_output_func("Application has been completed.\n")
    flag = messagebox.askquestion("Information", "Application has been completed.\nPlease check the download folder for the stock data analysis report.")
    if flag == "yes":
        os.system("start C:\\Users\\{}\\Downloads".format(username))

    else:
        pass

def Search():
    global ticker_list
    status_output_func("Searching...")
    URL = "https://finance.yahoo.com/"
    ticker = stc_sym.get()
    flag = check_ticker_exists(ticker)
    if flag:
        driver = webdriver.Chrome()
        driver.get(URL + "quote/{}".format(ticker))

        opt = webdriver.ChromeOptions()
        opt.add_experimental_option('detach', True) 
        opt.add_argument('--start-maximized')

    else:
        status_output_func("The {} does not exist in Yahoo Finance.\n".format(ticker))
        messagebox.showerror("Error", "The stock symbol does not exist in Yahoo Finance.")

    # if len(ticker_list) == 0:
    #     messagebox.showerror("Error", "No stock symbol has been added.")
    #     status_output_func("No stock symbol has been added.\n")

    # else:
    #     driver = webdriver.Chrome()
    #     for i in range(len(ticker_list)):
    #         ticker = ticker_list[i]
    #         driver.get(URL + "quote/{}".format(ticker))
    #         # time.sleep(0.5)
    #         pyautogui.hotkey("ctrl", "t")        
    #         # time.sleep(0.5)

def main():
    win = ttk.Window(themename="cerculean")
    win.title("Office Automation - Automatic Stock Data Analysis")
    win.geometry("1280x720")
    win.resizable(False, False)
    ico_path = fd.path_function("/NTTU_LOGO.ico")
    win.iconbitmap(ico_path)
    pic_bg_path = fd.path_function("/background.png")
    pic_bg = Image.open(pic_bg_path)
    pic = ImageTk.PhotoImage(pic_bg)
    tk.Label(win, image=pic).place(x=-2, y=-2)

    global stc_sym, date_from, date_to, status_output, rect_stock_output, PTC, VC, K, MAC, RSI, BOL, email_func, SMTP, email_add, password
    stc_sym = tk.StringVar()
    date_from = tk.StringVar()
    date_to = tk.StringVar()
    PTC = tk.BooleanVar()
    VC = tk.BooleanVar()
    K = tk.BooleanVar()
    MAC = tk.BooleanVar()
    RSI = tk.BooleanVar()
    BOL = tk.BooleanVar()
    email_func = tk.BooleanVar()
    SMTP = tk.StringVar()
    email_add = tk.StringVar()
    password = tk.StringVar()

    # date_from.set("YYYY/MM/DD")
    # date_to.set("YYYY/MM/DD")
    date_from.set((datetime.now() - timedelta(365)).strftime("%Y/%m/%d"))
    date_to.set(datetime.now().strftime("%Y/%m/%d"))
    PTC.set(True)
    SMTP.set("SMTP Server/Port")
    email_add.set("Email Address")
    password.set("Password")

    tk.Label(win, text="Symbol", font=("Arial", 20)).place(x=10, y=85)
    tk.Entry(win, font=("Arial", 20), textvariable=stc_sym).place(x=10, y=130, width=430, height=40)
    tk.Button(win, text="Apply", font=("Arial", 20), command=apply).place(x=450, y=130, width=180, height=40)

    tk.Label(win, text="Recently Searched", font=("Arial", 20)).place(x=10, y=190)
    rect_stock_output = tk.Text(win, font=("Arial", 20))
    rect_stock_output.place(x=10, y=235, width=620, height=175)
    rect_stock_output.config(state="disabled")
    tk.Button(win, text="Apply All", font=("Arial", 20), command=appl_all).place(x=10, y=420, width=620, height=40)

    print(rect_stock)
    rect_stock_output_func("\n".join(rect_stock))

    tk.Label(win, text="Date From", font=("Arial", 20)).place(x=640, y=85)
    tk.Entry(win, font=("Arial", 20), textvariable=date_from).place(x=640, y=130, width=310, height=40)
    tk.Label(win, text="Date To", font=("Arial", 20)).place(x=960, y=85)
    tk.Entry(win, font=("Arial", 20), textvariable=date_to).place(x=960, y=130, width=310, height=40)

    tk.Label(win, text="Functions", font=("Arial", 20)).place(x=640, y=190)
    # CB_PTC = tk.Checkbutton(win, text="Price Trend Chart", font=("Arial", 18), variable=PTC).place(x=640, y=235)
    # CB_VC = tk.Checkbutton(win, text="Volume Chart", font=("Arial", 18), variable=VC).place(x=640, y=270)
    # CB_K = tk.Checkbutton(win, text="K Line", font=("Arial", 18), variable=K).place(x=640, y=305)
    # CB_MAC = tk.Checkbutton(win, text="Moving Average Chart", font=("Arial", 18), variable=MAC).place(x=640, y=340)
    # CB_RSI = tk.Checkbutton(win, text="RSI Chart", font=("Arial", 18), variable=RSI).place(x=640, y=375)
    # CB_BOL = tk.Checkbutton(win, text="Bollinger Band Chart", font=("Arial", 18), variable=BOL).place(x=640, y=410)
    # CB_email_func = tk.Checkbutton(win, text="Email Notification", font=("Arial", 18), variable=email_func).place(x=640, y=445)
    CB_PTC = tk.Checkbutton(win, text="價格走勢圖", font=("Arial", 18), variable=PTC).place(x=640, y=235)
    CB_VC = tk.Checkbutton(win, text="成交量圖表", font=("Arial", 18), variable=VC).place(x=640, y=270)
    CB_K = tk.Checkbutton(win, text="K線圖", font=("Arial", 18), variable=K).place(x=640, y=305)
    CB_MAC = tk.Checkbutton(win, text="移動平均線圖", font=("Arial", 18), variable=MAC).place(x=640, y=340)
    CB_RSI = tk.Checkbutton(win, text="相對強弱指數圖表", font=("Arial", 18), variable=RSI).place(x=640, y=375)
    CB_BOL = tk.Checkbutton(win, text="布林線圖", font=("Arial", 18), variable=BOL).place(x=640, y=410)
    CB_email_func = tk.Checkbutton(win, text="Email Notification", font=("Arial", 18), variable=email_func).place(x=640, y=445)

    tk.Label(win, text="Mail Settings", font=("Arial", 20)).place(x=960, y=190)
    tk.Entry(win, font=("Arial", 20), textvariable=SMTP).place(x=960, y=235, width=310, height=40)
    tk.Entry(win, font=("Arial", 20), textvariable=email_add).place(x=960, y=285, width=310, height=40)
    tk.Entry(win, font=("Arial", 20), show='*', textvariable=password).place(x=960, y=335, width=310, height=40)
    tk.Button(win, text="Save to Default", font=("Arial", 20), command=save_to_default).place(x=960, y=390, width=310, height=40)
    tk.Button(win, text="Load Default", font=("Arial", 20), command=load_default).place(x=960, y=435, width=310, height=40)

    tk.Label(win, text="Status", font=("Arial", 20)).place(x=10, y=475)
    status_output = tk.Text(win, font=("Arial", 20))
    status_output.place(x=10, y=520, width=1260, height=100)
    status_output.config(state="disabled")

    tk.Button(win, text="Stock List", font=("Arial", 20), command=stock_list).place(x=10, y=635, width=200, height=40)
    tk.Button(win, text="Search - Yahoo Finance", font=("Arial", 20), command=Search).place(x=220, y=635, width=400, height=40)
    tk.Button(win, text="Run Application", font=("Arial", 20), command=application).place(x=630, y=635, width=640, height=40)

    win.mainloop()

if __name__ == "__main__":
    main()